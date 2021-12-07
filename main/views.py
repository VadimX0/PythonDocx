from .forms import TaskForm, DocForm
from .models import Task, chkboxcourse
from django.shortcuts import render, redirect
from docx import Document
from django.http import HttpResponse
from docxcompose.composer import Composer
from docx.shared import Inches, Pt
from django.views import View
import os

def index(request):
    tasks = Task.objects.order_by('id')
    return render(request, 'main/index.html', {'title': 'Главная страница сайта', 'tasks': tasks})


def about(request):
    return render(request, 'main/about.html')


def create(request): #Создание заданий
    error = ''
    if request.method == 'POST':
        form = TaskForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('home')
        else:
            error='Форма была неверной'

    form = TaskForm()
    context = {
        'form': form,
        'error': error
    }
    return render(request, 'main/create.html', context)


def generate_docx(request):#Функция создает файл в формате doc
    error = ''
    if request.method == 'POST':
        form = DocForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('home')
        else:
            error = 'Форма была неверной'

    form = DocForm()
    context = {
        'form': form,
        'error': error
    }
    return render(request, 'main/doc.html', context)
    return response





def savevalues(request):#Функция сохранения значений чекбоксов для страницы Чекбоксы
   if request.method == "POST":
      if request.POST.get('coursename'):
         savedata = chkboxcourse()
         savedata.coursename = request.POST.get('coursename')
         savedata.save()
         return render(request, 'main/doc.html')
   else:
         return render(request, 'main/doc.html')




class DownClass(View): #Класс для создания и загрузки докмента
  def __init__(self):
      self.first=Document(os.path.abspath('main/file1.docx'))
      self.second = Document(os.path.abspath('main/file2.docx'))
      self.docmx = Document(os.path.abspath('main/filex.docx')) #Основной документ
      self.files=[]
      super().__init__()
  def get(self, request):
       storyform = DocForm(request.GET)
       if storyform.is_valid():
        manname = storyform.cleaned_data['naming']
        namel = storyform.cleaned_data['name']
        org = storyform.cleaned_data['org']
        h = str(manname)#Имя человека
        n = str(namel)#Название документа
        o = str(org)#Название организации
        tab = self.docmx.add_table(rows=2, cols=1)
        cell1 = tab.cell(0, 0)#Поля, куда будут вводиться названия внутри документа
        cell2 = tab.cell(1, 0)
        cell1.text = "Вас зовут: " + h
        cell2.text = "Название организации: " + o
        self.docmx.save("doci.docx")
        merged_document = Document()
        self.files.append(self.docmx)
        for index, file in enumerate(self.files): #Создание документа взависимости от выбранного в чекбоксе количества файлов
           sub_doc = file
           if len(self.files)<=2:
             if index < len(self.files) - 1:
               sub_doc.add_page_break()
             for element in sub_doc.element.body:
               merged_document.element.body.append(element)
           elif len(self.files)>2:
                if index < len(self.files) - 2:
                    sub_doc.add_page_break()
                for element in sub_doc.element.body:
                    merged_document.element.body.append(element)
        style = merged_document.styles['Normal']
        font = style.font
        font.name = 'TimesNewRoman'
        font.size = Pt(14)
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=' + n + '.docx'
        merged_document.save(response)
        return response

  def download1(self, request): #соединить основной и первый файлы
      x = [self.docmx, self.first]
      self.files.extend(x)
      return self.get(request)

  def download2(self, request):#Все три
      x=[self.docmx, self.first, self.second]
      self.files.extend(x)
      return self.get(request)

  def download3(self, request):#Основной и второй
      x = [self.docmx,self.second]
      self.files.extend(x)
      return self.get(request)

def download11(request):
    return DownClass().download1(request)

def download22(request):
    return DownClass().download2(request)

def download33(request):
    return DownClass().download3(request)


